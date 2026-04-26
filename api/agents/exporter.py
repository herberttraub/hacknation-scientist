"""ExperimentPlan -> Markdown -> {pdf, docx, tex, md}.

PDF uses weasyprint (HTML->PDF, no LaTeX dependency).
DOCX and LaTeX use pandoc when available.
Plain markdown is always available.
"""
from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from html import escape
from pathlib import Path
from typing import Any


def _pandoc() -> str | None:
    found = shutil.which("pandoc")
    if found:
        return found
    candidates = [
        Path(os.environ.get("LOCALAPPDATA", "")) / "Pandoc" / "pandoc.exe",
        Path(os.environ.get("ProgramFiles", "")) / "Pandoc" / "pandoc.exe",
    ]
    for c in candidates:
        if c.exists():
            return str(c)
    return None


# ─── Markdown render ──────────────────────────────────────────────────────
def plan_to_markdown(plan: dict[str, Any]) -> str:
    lines: list[str] = []
    p = lambda s="": lines.append(s)

    p(f"# {plan.get('title', 'Experiment Plan')}")
    p()
    p(f"> *{plan.get('hypothesis', '')}*")
    p()
    p("## Novelty")
    p(plan.get("novelty_summary", ""))
    p()

    env = plan.get("environmental_conditions") or {}
    p("## Environmental Conditions")
    p(f"- **Temperature:** {env.get('temp_min_C')}–{env.get('temp_max_C')} °C")
    if env.get("humidity_min_pct") is not None:
        p(f"- **Humidity:** {env.get('humidity_min_pct')}–{env.get('humidity_max_pct')} %")
    if env.get("light"):
        p(f"- **Light:** {env['light']}")
    if env.get("atmosphere"):
        p(f"- **Atmosphere:** {env['atmosphere']}")
    if env.get("season_sensitivity"):
        p(f"- **Season sensitivity:** {env['season_sensitivity']}")
    p()

    p("## Protocol")
    for i, s in enumerate(plan.get("protocol") or [], 1):
        parallel = ", ".join(s.get("can_run_parallel_with") or [])
        p(f"### Step {i} — {s.get('name')}  *({s.get('duration_minutes', 0)} min)*")
        p(f"**Rationale.** {s.get('rationale', '')}")
        if s.get("materials_used"):
            p(f"- Materials: {', '.join(s['materials_used'])}")
        if s.get("equipment_used"):
            p(f"- Equipment: {', '.join(s['equipment_used'])}")
        if s.get("qc_checks"):
            p(f"- QC: {', '.join(s['qc_checks'])}")
        if s.get("assumed_skills"):
            p(f"- Assumes: {', '.join(s['assumed_skills'])}")
        if parallel:
            p(f"- Can run in parallel with: {parallel}")
        if s.get("notes"):
            p(f"> {s['notes']}")
        p()

    p("## Materials")
    for m in plan.get("materials") or []:
        shelf = f"; shelf life {m.get('shelf_life_days')} days" if m.get("shelf_life_days") else ""
        catalog = f", catalog {m.get('catalog_no')}" if m.get("catalog_no") else ""
        supplier = f", supplier {m.get('supplier')}" if m.get("supplier") else ""
        order = f", order {m.get('order_priority')}" if m.get("order_priority") else ""
        storage = f", storage {m.get('storage')}{shelf}" if m.get("storage") or shelf else ""
        p(
            f"- **{m.get('name','')}**{catalog}{supplier}: "
            f"{m.get('qty','')} {m.get('unit_size','')} at "
            f"${m.get('unit_cost_usd', 0):.2f} each; total ${m.get('total_cost_usd', 0):.2f}"
            f"{order}{storage}."
        )
    p()

    p("## Equipment")
    for e in plan.get("equipment") or []:
        p(f"- **{e.get('name','')}** ({e.get('model','')}) — {e.get('location','')}")
    p()

    b = plan.get("budget") or {}
    p("## Budget")
    p(f"**Total: ${b.get('total_usd', 0):,.0f} {b.get('currency','USD')}** · {b.get('contingency_pct', 0):.0f}% contingency")
    p()
    if b.get("categories"):
        for c in b["categories"]:
            p(f"- {c.get('name')}: ${c.get('total_usd', 0):,.0f}")
        p()
    if plan.get("budget_justification"):
        p(f"> {plan['budget_justification']}")
        p()

    p("## Timeline")
    for ph in plan.get("timeline") or []:
        deps = f"  *(depends on: {', '.join(ph.get('dependencies') or [])})*" if ph.get("dependencies") else ""
        p(f"- **w{ph.get('week_start')}–w{ph.get('week_end')} {ph.get('name','')}**{deps}")
        for d in ph.get("deliverables") or []:
            p(f"  - {d}")
    p()

    p("## Staffing")
    for s in plan.get("staffing") or []:
        p(f"- **{s.get('role','')}** — {s.get('named_person','')} ({s.get('institution','')}) · {s.get('fte_pct',0)}% FTE")
        if s.get("expertise_tags"):
            p(f"  - Expertise: {', '.join(s['expertise_tags'])}")
    p()

    v = plan.get("validation") or {}
    p("## Validation")
    p()
    p("**Success criteria:**")
    for c in v.get("success_criteria") or []:
        p(f"- {c}")
    if v.get("failure_modes"):
        p()
        p("**Failure modes:**")
        for f in v["failure_modes"]:
            p(f"- *{f}*")
    if v.get("statistics_plan"):
        p()
        p(f"**Statistics.** {v['statistics_plan']}")
    p()

    if plan.get("open_questions"):
        p("## Open questions for the scientist")
        for q in plan["open_questions"]:
            p(f"- {q}")
        p()

    if plan.get("references"):
        p("## References")
        for r in plan["references"]:
            line = f"- *{r.get('title','')}* — {r.get('authors','')}"
            if r.get("year"):
                line += f" ({r['year']})"
            if r.get("doi") or r.get("url"):
                line += f"  [{r.get('doi') or r.get('url')}]({r.get('url') or 'https://doi.org/' + (r.get('doi') or '')})"
            p(line)
        p()

    return "\n".join(lines)


# ─── HTML for PDF (weasyprint) ────────────────────────────────────────────
PDF_CSS = """
@page { size: Letter; margin: 0.9in; }
body { font-family: 'EB Garamond', 'Garamond', Georgia, serif; color: #2B2B2B; line-height: 1.5; font-size: 11pt; }
h1 { font-size: 26pt; margin: 0 0 8pt; border-bottom: 1px solid #A8794A; padding-bottom: 6pt; }
h2 { font-size: 16pt; margin: 22pt 0 6pt; color: #2B2B2B; border-bottom: 1px solid #D9CFBE; padding-bottom: 3pt; }
h3 { font-size: 13pt; margin: 14pt 0 4pt; font-style: italic; }
blockquote { border-left: 2px solid #A8794A; padding-left: 8pt; color: #5a5446; font-style: italic; margin: 10pt 0; }
code, pre { font-family: 'JetBrains Mono', monospace; font-size: 9pt; }
table { width: 100%; border-collapse: collapse; margin: 8pt 0; font-size: 8pt; table-layout: fixed; }
th, td { padding: 3pt 4pt; border-bottom: 0.5pt solid #D9CFBE; text-align: left; word-wrap: break-word; overflow-wrap: break-word; }
th { background: #F4EFE6; color: #A8794A; text-transform: uppercase; letter-spacing: 0.1em; font-size: 8pt; font-weight: 500; }
ul, ol { padding-left: 18pt; }
li { margin: 2pt 0; }
strong { color: #2B2B2B; }
em { color: #5a5446; }
p, li, blockquote { page-break-inside: avoid; }
"""

PRETTY_PDF_CSS = """
@page { size: Letter; margin: 0.55in; }
body { font-family: Georgia, 'Times New Roman', serif; color: #2B2B2B; line-height: 1.42; font-size: 10pt; }
.eyebrow { font-family: Helvetica, Arial, sans-serif; text-transform: uppercase; letter-spacing: 1pt; font-size: 7pt; color: #A8794A; font-weight: bold; }
.cover { border-bottom: 1.5pt solid #A8794A; padding-bottom: 12pt; margin-bottom: 14pt; }
h1 { font-size: 24pt; line-height: 1.05; margin: 4pt 0 8pt; font-weight: normal; }
h2 { font-size: 15pt; margin: 18pt 0 6pt; border-bottom: 0.75pt solid #A8794A; padding-bottom: 3pt; font-weight: normal; }
h3 { font-size: 12pt; margin: 7pt 0 3pt; font-weight: normal; }
p { margin: 4pt 0 7pt; }
.meta { color: #6b6254; font-style: italic; max-width: 94%; }
.grid { display: table; width: 100%; table-layout: fixed; border-spacing: 7pt; margin: 8pt -7pt; }
.cell { display: table-cell; border: 0.75pt solid #D9CFBE; padding: 8pt; vertical-align: top; background: #FBF7EF; }
.stat { font-family: Courier, monospace; font-size: 16pt; margin-top: 4pt; }
.bar-row { margin: 5pt 0; page-break-inside: avoid; }
.bar-label { display: inline-block; width: 38%; font-size: 8pt; vertical-align: middle; }
.bar-track { display: inline-block; width: 40%; height: 8pt; background: #E6DCCB; vertical-align: middle; }
.bar-fill { height: 8pt; background: #A8794A; }
.bar-value { display: inline-block; width: 18%; text-align: right; font-family: Courier, monospace; font-size: 8pt; }
.phase { border-left: 3pt solid #A8794A; padding: 5pt 7pt; margin: 5pt 0; background: #FBF7EF; page-break-inside: avoid; }
.step { border-bottom: 0.5pt solid #D9CFBE; padding: 5pt 0; page-break-inside: avoid; }
.chip { font-family: Helvetica, Arial, sans-serif; font-size: 7pt; color: #A8794A; }
table { width: 100%; border-collapse: collapse; margin: 7pt 0; font-size: 8pt; table-layout: fixed; }
th, td { padding: 4pt; border-bottom: 0.5pt solid #D9CFBE; text-align: left; word-wrap: break-word; }
th { font-family: Helvetica, Arial, sans-serif; color: #A8794A; text-transform: uppercase; letter-spacing: 0.75pt; font-size: 7pt; }
ul { padding-left: 13pt; margin: 4pt 0; }
li { margin: 2pt 0; }
.quadrant { display: table; width: 100%; table-layout: fixed; border-spacing: 7pt; margin: 8pt -7pt; }
.page-break { page-break-before: always; }
"""


def _money(value: Any) -> str:
    try:
        return f"${float(value):,.0f}"
    except Exception:
        return "$0"


def _list(items: list[str]) -> str:
    return "<ul>" + "".join(f"<li>{escape(str(item))}</li>" for item in items) + "</ul>"


def _budget_bars(materials: list[dict[str, Any]]) -> str:
    rows = sorted(materials, key=lambda m: float(m.get("total_cost_usd") or 0), reverse=True)
    max_cost = max([float(m.get("total_cost_usd") or 0) for m in rows] or [1])
    out = []
    for m in rows:
        cost = float(m.get("total_cost_usd") or 0)
        width = max(3, int((cost / max_cost) * 100)) if max_cost else 3
        out.append(
            "<div class='bar-row'>"
            f"<span class='bar-label'>{escape(m.get('name', 'Material'))}</span>"
            f"<span class='bar-track'><span class='bar-fill' style='width:{width}%;'></span></span>"
            f"<span class='bar-value'>{_money(cost)}</span>"
            "</div>"
        )
    return "".join(out)


def plan_to_pretty_html(plan: dict[str, Any]) -> str:
    title = escape(plan.get("title") or "Experiment Plan")
    hypothesis = escape(plan.get("hypothesis") or "")
    env = plan.get("environmental_conditions") or {}
    budget = plan.get("budget") or {}
    protocol = plan.get("protocol") or []
    materials = plan.get("materials") or []
    equipment = plan.get("equipment") or []
    timeline = plan.get("timeline") or []
    validation = plan.get("validation") or {}
    refs = plan.get("references") or []
    total_weeks = max([int(p.get("week_end") or 0) for p in timeline] or [0])

    material_rows = "".join(
        "<tr>"
        f"<td>{escape(m.get('name', ''))}</td>"
        f"<td>{escape(m.get('supplier', ''))}</td>"
        f"<td>{escape(m.get('catalog_no', ''))}</td>"
        f"<td>{escape(str(m.get('qty', '')))} {escape(m.get('unit_size', ''))}</td>"
        f"<td>{_money(m.get('total_cost_usd') or 0)}</td>"
        "</tr>"
        for m in materials
    )
    equipment_rows = "".join(
        "<tr>"
        f"<td>{escape(e.get('name', ''))}</td>"
        f"<td>{escape(e.get('model', ''))}</td>"
        f"<td>{escape(e.get('location', ''))}</td>"
        f"<td>{escape(e.get('owner_team', ''))}</td>"
        "</tr>"
        for e in equipment
    )
    steps = "".join(
        "<div class='step'>"
        f"<div class='eyebrow'>Step {i} · {int(s.get('duration_minutes') or 0)} min</div>"
        f"<h3>{escape(s.get('name', ''))}</h3>"
        f"<p>{escape(s.get('rationale', ''))}</p>"
        f"<div class='chip'>Materials: {escape(', '.join(s.get('materials_used') or []) or 'none listed')}</div>"
        f"<div class='chip'>Equipment: {escape(', '.join(s.get('equipment_used') or []) or 'none listed')}</div>"
        "</div>"
        for i, s in enumerate(protocol, 1)
    )
    phases = "".join(
        "<div class='phase'>"
        f"<div class='eyebrow'>Week {p.get('week_start')} to {p.get('week_end')}</div>"
        f"<h3>{escape(p.get('name', ''))}</h3>"
        f"{_list(p.get('deliverables') or []) if p.get('deliverables') else ''}"
        "</div>"
        for p in timeline
    )
    refs_html = "".join(
        f"<li><em>{escape(r.get('title', ''))}</em> {escape(r.get('authors', ''))} {escape(str(r.get('year') or ''))}</li>"
        for r in refs
    )

    body = f"""
    <div class="cover">
      <div class="eyebrow">Husky Lab · AI Scientist</div>
      <h1>{title}</h1>
      <p class="meta">{hypothesis}</p>
    </div>

    <div class="grid">
      <div class="cell"><div class="eyebrow">Duration</div><div class="stat">{total_weeks} weeks</div></div>
      <div class="cell"><div class="eyebrow">Protocol</div><div class="stat">{len(protocol)} steps</div></div>
      <div class="cell"><div class="eyebrow">Budget</div><div class="stat">{_money(budget.get('total_usd') or 0)}</div></div>
    </div>

    <h2>Overview</h2>
    <p>{escape(plan.get('novelty_summary') or '')}</p>
    <div class="grid">
      <div class="cell"><div class="eyebrow">Temperature</div><div class="stat">{env.get('temp_min_C', '')}-{env.get('temp_max_C', '')}C</div></div>
      <div class="cell"><div class="eyebrow">Humidity</div><div class="stat">{env.get('humidity_min_pct', '')}-{env.get('humidity_max_pct', '')}%</div></div>
      <div class="cell"><div class="eyebrow">Atmosphere</div><p>{escape(env.get('atmosphere') or 'standard bench conditions')}</p></div>
    </div>

    <h2>Timeline</h2>
    {phases}

    <h2>Budget Visualization</h2>
    {_budget_bars(materials)}

    <h2 class="page-break">Protocol</h2>
    {steps}

    <h2>Materials</h2>
    <table><thead><tr><th>Material</th><th>Supplier</th><th>Catalog</th><th>Qty</th><th>Total</th></tr></thead><tbody>{material_rows}</tbody></table>

    <h2>Equipment</h2>
    <table><thead><tr><th>Equipment</th><th>Model</th><th>Location</th><th>Owner</th></tr></thead><tbody>{equipment_rows}</tbody></table>

    <h2>Validation</h2>
    <div class="quadrant">
      <div class="cell"><div class="eyebrow">Success</div>{_list(validation.get('success_criteria') or [])}</div>
      <div class="cell"><div class="eyebrow">Failure Modes</div>{_list(validation.get('failure_modes') or [])}</div>
    </div>
    <div class="cell"><div class="eyebrow">Statistics</div><p>{escape(validation.get('statistics_plan') or '')}</p></div>

    <h2>References</h2>
    <ol>{refs_html}</ol>
    """
    return f"<!DOCTYPE html><html><head><meta charset='utf-8'><style>{PRETTY_PDF_CSS}</style></head><body>{body}</body></html>"


def markdown_to_html(md_text: str) -> str:
    import markdown as md_mod

    body = md_mod.markdown(md_text, extensions=["tables", "fenced_code"])
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8"><style>{PDF_CSS}</style></head><body>{body}</body></html>"""


def render_pdf(md_text: str) -> bytes:
    """Use xhtml2pdf — pure Python, no system deps."""
    import io

    from xhtml2pdf import pisa

    html = markdown_to_html(md_text)
    buf = io.BytesIO()
    pisa.CreatePDF(html, dest=buf, encoding="utf-8")
    return buf.getvalue()


def render_plan_pdf(plan: dict[str, Any]) -> bytes:
    import io

    from xhtml2pdf import pisa

    buf = io.BytesIO()
    pisa.CreatePDF(plan_to_pretty_html(plan), dest=buf, encoding="utf-8")
    return buf.getvalue()


def render_docx(md_text: str) -> bytes:
    """Use pandoc if available, else fall back to python-docx (less pretty)."""
    pandoc = _pandoc()
    if pandoc:
        with tempfile.TemporaryDirectory() as td:
            tdp = Path(td)
            (tdp / "in.md").write_text(md_text, encoding="utf-8")
            out = tdp / "out.docx"
            subprocess.run(
                [pandoc, str(tdp / "in.md"), "-o", str(out), "--standalone"],
                check=True,
                capture_output=True,
            )
            return out.read_bytes()
    # fallback
    from docx import Document

    doc = Document()
    for line in md_text.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    buf = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(buf.name)
    return Path(buf.name).read_bytes()


def render_latex(md_text: str) -> str:
    pandoc = _pandoc()
    if pandoc:
        with tempfile.TemporaryDirectory() as td:
            tdp = Path(td)
            (tdp / "in.md").write_text(md_text, encoding="utf-8")
            out = tdp / "out.tex"
            subprocess.run(
                [pandoc, str(tdp / "in.md"), "-o", str(out), "--standalone"],
                check=True,
                capture_output=True,
            )
            return out.read_text(encoding="utf-8")
    # fallback: minimal hand-rolled
    import re

    tex = md_text
    tex = re.sub(r"^# (.+)$", r"\\section*{\1}", tex, flags=re.M)
    tex = re.sub(r"^## (.+)$", r"\\subsection*{\1}", tex, flags=re.M)
    tex = re.sub(r"^### (.+)$", r"\\subsubsection*{\1}", tex, flags=re.M)
    tex = re.sub(r"\*\*(.+?)\*\*", r"\\textbf{\1}", tex)
    tex = re.sub(r"\*(.+?)\*", r"\\textit{\1}", tex)
    return "\\documentclass[11pt]{article}\n\\usepackage{geometry}\n\\geometry{letterpaper, margin=1in}\n\\begin{document}\n" + tex + "\n\\end{document}\n"
